import os
import requests
import fitz
from docx import Document
from sqlalchemy.orm import Session
from fastapi import UploadFile

from app.db import crud


def save_document_links(db: Session, notice_id: str, document_links: list[dict]) -> int:
    saved = 0

    for link in document_links or []:
        url = link.get("url")
        if not url:
            continue

        crud.upsert_opportunity_document(
            db,
            {
                "notice_id": notice_id,
                "document_url": url,
                "document_name": link.get("name", "") or url.split("/")[-1],
                "document_type": link.get("type", "unknown"),
                "extraction_status": "pending",
            },
        )
        saved += 1

    return saved


def infer_file_type(name_or_url: str, content_type: str = "") -> str:
    value = (name_or_url or "").lower()
    content_type = (content_type or "").lower()

    if value.endswith(".pdf") or "pdf" in content_type:
        return "pdf"

    if value.endswith(".docx") or "wordprocessingml" in content_type:
        return "docx"

    if value.endswith(".txt") or "text/plain" in content_type:
        return "txt"

    return "unknown"


def extract_pdf_text_from_bytes(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = [page.get_text() for page in doc]
    return "\n".join(text_parts).strip()


def extract_docx_text_from_bytes(file_bytes: bytes) -> str:
    import io

    file_stream = io.BytesIO(file_bytes)
    document = Document(file_stream)

    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts).strip()


def extract_txt_text_from_bytes(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore").strip()


def extract_text_from_bytes(file_bytes: bytes, filename: str = "", content_type: str = "") -> tuple[str, str]:
    file_type = infer_file_type(filename, content_type)

    if file_type == "pdf":
        return extract_pdf_text_from_bytes(file_bytes), "pdf"

    if file_type == "docx":
        return extract_docx_text_from_bytes(file_bytes), "docx"

    if file_type == "txt":
        return extract_txt_text_from_bytes(file_bytes), "txt"

    raise ValueError(f"Unsupported or unknown document type: {filename} | {content_type}")


def extract_text_from_url(url: str) -> tuple[str, str]:
    response = requests.get(url, timeout=60)
    response.raise_for_status()

    content_type = response.headers.get("content-type", "")
    filename = url.split("?")[0].split("/")[-1]

    return extract_text_from_bytes(
        file_bytes=response.content,
        filename=filename,
        content_type=content_type,
    )


def process_document_rows(db: Session, docs) -> dict:
    processed = 0
    failed = 0

    for doc in docs:
        try:
            extracted_text, detected_type = extract_text_from_url(doc.document_url)

            crud.upsert_opportunity_document(
                db,
                {
                    "notice_id": doc.notice_id,
                    "document_url": doc.document_url,
                    "document_name": doc.document_name,
                    "document_type": detected_type,
                    "extracted_text": extracted_text,
                    "extraction_status": "complete",
                    "error_message": "",
                },
            )

            processed += 1

        except Exception as exc:
            crud.upsert_opportunity_document(
                db,
                {
                    "notice_id": doc.notice_id,
                    "document_url": doc.document_url,
                    "document_name": doc.document_name,
                    "document_type": doc.document_type,
                    "extracted_text": "",
                    "extraction_status": "failed",
                    "error_message": str(exc),
                },
            )

            failed += 1

    return {
        "processed": processed,
        "failed": failed,
        "attempted": len(docs),
    }


def process_pending_documents(db: Session, limit: int = 10) -> dict:
    docs = crud.get_pending_documents(db, limit=limit)
    return process_document_rows(db, docs)


def process_documents_for_reviewed_opportunities(db: Session, profile_id: int, limit: int = 10) -> dict:
    reviews = crud.get_reviews_for_profile(db, profile_id)

    target_notice_ids = [
        review.notice_id
        for review in reviews
        if review.disposition in ["Good Fit", "Maybe"]
    ]

    docs = crud.get_pending_documents_for_notice_ids(
        db=db,
        notice_ids=target_notice_ids,
        limit=limit,
    )

    result = process_document_rows(db, docs)
    result["reviewed_notice_ids"] = target_notice_ids
    return result


def upload_and_extract_document(db: Session, notice_id: str, file: UploadFile) -> dict:
    try:
        file_bytes = file.file.read()

        extracted_text, detected_type = extract_text_from_bytes(
            file_bytes=file_bytes,
            filename=file.filename,
            content_type=file.content_type or "",
        )

        document_url = f"manual_upload://{notice_id}/{file.filename}"

        doc = crud.upsert_opportunity_document(
            db,
            {
                "notice_id": notice_id,
                "document_url": document_url,
                "document_name": file.filename,
                "document_type": f"manual_{detected_type}",
                "extracted_text": extracted_text,
                "extraction_status": "complete",
                "error_message": "",
            },
        )

        return {
            "status": "uploaded_and_extracted",
            "notice_id": notice_id,
            "document_name": file.filename,
            "document_type": detected_type,
            "text_length": len(extracted_text),
            "document_id": doc.id,
        }

    except Exception as exc:
        document_url = f"manual_upload://{notice_id}/{file.filename}"

        crud.upsert_opportunity_document(
            db,
            {
                "notice_id": notice_id,
                "document_url": document_url,
                "document_name": file.filename,
                "document_type": "manual_unknown",
                "extracted_text": "",
                "extraction_status": "failed",
                "error_message": str(exc),
            },
        )

        return {
            "status": "failed",
            "notice_id": notice_id,
            "document_name": file.filename,
            "error": str(exc),
        }


def get_documents_for_notice(db: Session, notice_id: str) -> list[dict]:
    docs = crud.get_documents_for_notice(db, notice_id)

    return [
        {
            "notice_id": d.notice_id,
            "document_url": d.document_url,
            "document_name": d.document_name,
            "document_type": d.document_type,
            "extraction_status": d.extraction_status,
            "error_message": d.error_message,
            "text_preview": (d.extracted_text or "")[:1000],
            "text_length": len(d.extracted_text or ""),
        }
        for d in docs
    ]


def get_combined_document_text(db: Session, notice_id: str, max_chars: int = 12000) -> str:
    docs = crud.get_documents_for_notice(db, notice_id)

    text = ""

    for doc in docs:
        if doc.extraction_status == "complete" and doc.extracted_text:
            text += f"\n\nDOCUMENT: {doc.document_name}\n"
            text += doc.extracted_text

    return text[:max_chars]