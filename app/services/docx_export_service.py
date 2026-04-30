from io import BytesIO
from docx import Document


def build_revised_proposal_docx(revised_proposal: dict, review: dict | None = None) -> BytesIO:
    document = Document()

    document.add_heading("Revised Proposal Draft", level=1)

    sections = [
        ("Executive Summary", "executive_summary"),
        ("Technical Approach", "technical_approach"),
        ("Management Plan", "management_plan"),
        ("Past Performance", "past_performance"),
        ("Staffing Plan", "staffing_plan"),
    ]

    for heading, key in sections:
        document.add_heading(heading, level=2)
        document.add_paragraph(revised_proposal.get(key, ""))

    if review:
        document.add_page_break()
        document.add_heading("Senior Capture Advisor Review Notes", level=1)

        document.add_heading("Overall Assessment", level=2)
        document.add_paragraph(review.get("overall_assessment", ""))

        document.add_heading("Screening Risks", level=2)
        for item in review.get("screening_risks", []):
            document.add_paragraph(item, style="List Bullet")

        document.add_heading("Compliance Gaps", level=2)
        for item in review.get("compliance_gaps", []):
            document.add_paragraph(item, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer